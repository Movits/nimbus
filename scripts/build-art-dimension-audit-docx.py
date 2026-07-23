from __future__ import annotations

import csv
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(r"C:\Users\rober\Nimbus")
AUDIT_DIR = ROOT / "nuvemshop" / "auditoria" / "2026-07-22-dimensoes-arte"
CSV_PATH = AUDIT_DIR / "auditoria-dimensoes-arte.csv"
CARDS_DIR = AUDIT_DIR / "cards"
CARDS_DOCX_DIR = AUDIT_DIR / "cards-docx"
LOGO_PATH = ROOT / "nuvemshop" / "assets" / "logo-nimbus.png"
OUTPUT_PATH = AUDIT_DIR / "auditoria-dimensoes-arte-nimbus.docx"

# compact_reference_guide preset + editorial_cover pattern.
# Landscape is a named override because the core deliverable is a side-by-side visual audit.
NAVY = "10265E"
LIGHT_BLUE = "D7ECFF"
PALE_BLUE = "EDF6FF"
GOLD = "E3B63F"
GREEN = "20674F"
PALE_GREEN = "E9F7F1"
RED = "A91E22"
PALE_RED = "FDEEEE"
INK = "18315B"
MUTED = "476387"
WHITE = "FFFFFF"


def read_rows() -> list[dict[str, str]]:
    with CSV_PATH.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, *, size=9, bold=False, color=INK, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("NIMBUS  |  Auditoria visual  •  ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.18)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(NAVY)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(10)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(NAVY)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(MUTED)
    h3.paragraph_format.space_before = Pt(4)
    h3.paragraph_format.space_after = Pt(3)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_small_caps(doc: Document, text: str, *, color=MUTED, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.character_spacing = Pt(0.6)


def add_cover(doc: Document) -> None:
    banner = doc.add_table(rows=1, cols=2)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = False
    banner.columns[0].width = Inches(6.55)
    banner.columns[1].width = Inches(3.55)
    left, right = banner.rows[0].cells
    for cell in (left, right):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=260, start=300, bottom=260, end=300)
    set_cell_shading(left, LIGHT_BLUE)
    set_cell_shading(right, NAVY)

    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("AUDITORIA VISUAL DE PRODUTOS")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Escala real das estampas")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(32)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("YouDraw × fotos de modelo publicadas na Nuvemshop")
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(INK)

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("49 produtos  •  49 comparações lado a lado  •  medidas oficiais em centímetros")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        rp.add_run().add_picture(str(LOGO_PATH), width=Inches(2.25))
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp2.paragraph_format.space_before = Pt(16)
    rr = rp2.add_run("22 JUL 2026")
    rr.bold = True
    rr.font.name = "Calibri"
    rr.font.size = Pt(11)
    rr.font.color.rgb = RGBColor.from_string(GOLD)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    note = doc.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = note.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, top=140, start=220, bottom=140, end=220)
    set_cell_text(
        cell,
        "Objetivo: identificar quais capas podem permanecer, quais exigem revisão e quais precisam ser refeitas antes de investir novos créditos em geração de imagem.",
        size=11,
        color=INK,
    )


def add_stat_box(cell, number: int, label: str, fill: str, accent: str) -> None:
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=170, bottom=150, end=170)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(str(number))
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(accent)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(label)
    r2.bold = True
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(accent)


def add_executive_summary(doc: Document, rows: list[dict[str, str]]) -> None:
    doc.add_heading("Resultado executivo", level=1)
    p = doc.add_paragraph(
        "A auditoria cruzou as dimensões oficiais informadas pela YouDraw com a proporção visual da estampa nas capas atuais da loja. O resultado não autoriza alterações automáticas: ele organiza o seu feedback antes de qualquer nova geração ou publicação."
    )
    p.paragraph_format.space_after = Pt(12)

    counts = Counter(row["verdict"] for row in rows)
    stats = doc.add_table(rows=1, cols=3)
    stats.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats.autofit = False
    for i, width in enumerate((3.25, 3.25, 3.25)):
        stats.columns[i].width = Inches(width)
    add_stat_box(stats.cell(0, 0), counts["APROVAR"], "APROVAR", PALE_GREEN, GREEN)
    add_stat_box(stats.cell(0, 1), counts["REVISAR"], "REVISAR", "FFF5D9", "8C6500")
    add_stat_box(stats.cell(0, 2), counts["REFAZER"], "REFAZER", PALE_RED, RED)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    doc.add_heading("Leitura mais importante", level=2)
    bullets = [
        "Os mockups e as medidas da YouDraw são a fonte de verdade; a foto de modelo nunca deve redesenhar a arte.",
        "O maior risco atual é a inconsistência entre peças da mesma família, não apenas um erro isolado de tamanho.",
        "São Jorge Neobarroco, São Miguel Vitorioso, São Jorge Vintage e Anjo da Guarda Stencil concentram as correções mais visíveis.",
        "A estimativa percentual vem da proporção no painel útil da peça; perspectiva, pose e caimento impedem converter a foto em centímetros exatos.",
    ]
    for text in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(text)


def add_methodology(doc: Document) -> None:
    doc.add_heading("Método e limites da auditoria", level=1)
    methods = [
        ("1. Fonte de verdade", "Dimensões exatas de frente e costas coletadas dentro de cada produto no painel autenticado da YouDraw."),
        ("2. Comparação visual", "Foto de modelo atual comparada lado a lado com os mockups oficiais de todas as cores disponíveis."),
        ("3. Leitura de escala", "Comparação da área ocupada pela arte em relação ao painel útil da frente ou das costas, compensando visualmente perspectiva e caimento."),
        ("4. Identidade", "Conferência de composição, moldura, personagens, tipografia, microtexto e cores. Arte apenas parecida não é considerada fiel."),
        ("5. Consistência", "Quando há mais de uma cor ou peça da mesma arte, a escala aparente também precisa permanecer estável."),
    ]
    table = doc.add_table(rows=len(methods), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.25)
    table.columns[1].width = Inches(7.45)
    for idx, (label, body) in enumerate(methods):
        c1, c2 = table.rows[idx].cells
        fill = LIGHT_BLUE if idx % 2 == 0 else PALE_BLUE
        set_cell_shading(c1, fill)
        set_cell_shading(c2, fill)
        set_cell_margins(c1)
        set_cell_margins(c2)
        set_cell_text(c1, label, size=10, bold=True, color=NAVY)
        set_cell_text(c2, body, size=10, color=INK)

    doc.add_heading("Faixas de decisão", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    for label, text, color in (
        ("APROVAR", " escala aparente coerente, geralmente dentro de aproximadamente ±8–10%, sem alteração estrutural da arte.", GREEN),
        ("REVISAR", " desvio limítrofe, baixa confiança em microdetalhes ou dúvida de caimento/modelagem.", "8C6500"),
        ("REFAZER", " desvio visível, normalmente acima de 15%, inconsistência entre cores ou risco de arte regenerada.", RED),
    ):
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(color)
        p.add_run(text)
        p.add_run("\n")

    warning = doc.add_table(rows=1, cols=1)
    cell = warning.cell(0, 0)
    set_cell_shading(cell, "FFF7E1")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    set_cell_text(
        cell,
        "Limite técnico: uma fotografia em perspectiva não permite provar centímetros exatos. Por isso, o relatório separa medida oficial (exata) de leitura fotográfica (estimada e acompanhada por confiança).",
        size=10,
        bold=True,
        color="745500",
    )


def priority_label(row: dict[str, str]) -> str:
    name = row["title"].replace(" | ", " — ")
    return f"{name}: {row['recommendation']}"


def add_priority_page(doc: Document, rows: list[dict[str, str]], verdict: str, title: str, intro: str) -> None:
    doc.add_heading(title, level=1)
    p = doc.add_paragraph(intro)
    p.paragraph_format.space_after = Pt(8)
    subset = [row for row in rows if row["verdict"] == verdict]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.85)
    table.columns[1].width = Inches(4.85)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, ("Produto / diagnóstico", "Produto / diagnóstico")):
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        set_cell_text(cell, text, size=9, bold=True, color=WHITE)

    pairs = [subset[i : i + 2] for i in range(0, len(subset), 2)]
    for row_index, pair in enumerate(pairs):
        cells = table.add_row().cells
        fill = PALE_RED if verdict == "REFAZER" else ("FFF8E7" if row_index % 2 == 0 else "FFFBF1")
        for idx, cell in enumerate(cells):
            set_cell_shading(cell, fill)
            set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
            if idx < len(pair):
                entry = pair[idx]
                set_cell_text(cell, priority_label(entry), size=8.5, bold=False, color=INK)
            else:
                set_cell_text(cell, "", size=8.5)


def add_card_page(doc: Document, row: dict[str, str], index: int, total: int) -> None:
    card_path = CARDS_DIR / f"{row['product_id']}.jpg"
    if not card_path.exists():
        raise FileNotFoundError(card_path)
    # python-docx rejects some valid Sharp JPEG encodings. Re-encoding as a
    # baseline JPEG keeps the card visually equivalent and portable in Word.
    CARDS_DOCX_DIR.mkdir(parents=True, exist_ok=True)
    docx_card_path = CARDS_DOCX_DIR / f"{row['product_id']}.jpg"
    with Image.open(card_path).convert("RGB") as image:
        image = image.resize((1200, 700), Image.Resampling.LANCZOS)
        image.save(
            docx_card_path,
            format="JPEG",
            quality=88,
            subsampling=2,
            optimize=False,
            progressive=False,
        )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"COMPARAÇÃO {index:02d}/{total:02d}  •  {row['collection']}  •  {row['verdict']}")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    img = doc.add_paragraph()
    img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img.paragraph_format.space_before = Pt(0)
    img.paragraph_format.space_after = Pt(0)
    img.add_run().add_picture(str(docx_card_path), width=Inches(10.0))


def add_next_steps(doc: Document) -> None:
    doc.add_heading("Próxima etapa após o seu feedback", level=1)
    doc.add_paragraph(
        "Nenhuma capa deve ser substituída antes de você aprovar os diagnósticos deste relatório. Depois da aprovação, a correção deve ser feita em lotes pequenos e validada no site real."
    )
    steps = [
        ("1", "Priorizar os 11 itens REFAZER", "Começar pelas diferenças mais visíveis e pelas famílias inconsistentes."),
        ("2", "Usar a arte original como camada rígida", "Gerar apenas pessoa, cenário, luz e tecido; aplicar a estampa original sem redesenho generativo."),
        ("3", "Testar uma peça piloto", "Validar primeiro no Nano Banana/Gemini com um único produto e uma única cor."),
        ("4", "Aprovar escala e identidade", "Comparar novamente com o mockup YouDraw e registrar o delta exato."),
        ("5", "Escalar somente o método aprovado", "Usar Higgsfield apenas quando o prompt, a referência e a escala já estiverem comprovados."),
        ("6", "Publicar sem apagar o acervo", "Preservar todas as imagens oficiais da YouDraw e trocar apenas a capa lifestyle reprovada."),
    ]
    table = doc.add_table(rows=len(steps), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.65)
    table.columns[1].width = Inches(3.25)
    table.columns[2].width = Inches(5.8)
    for idx, (num, title, body) in enumerate(steps):
        cells = table.rows[idx].cells
        fill = LIGHT_BLUE if idx % 2 == 0 else PALE_BLUE
        for cell in cells:
            set_cell_shading(cell, fill)
            set_cell_margins(cell)
        set_cell_text(cells[0], num, size=14, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], title, size=10, bold=True, color=NAVY)
        set_cell_text(cells[2], body, size=10, color=INK)

    doc.add_heading("Regra de parada", level=2)
    rule = doc.add_table(rows=1, cols=1)
    cell = rule.cell(0, 0)
    set_cell_shading(cell, PALE_RED)
    set_cell_margins(cell, top=140, start=170, bottom=140, end=170)
    set_cell_text(
        cell,
        "Se duas tentativas repetirem o mesmo erro de identidade ou escala, parar de gastar créditos e mudar o método — preferencialmente para composição com máscara/camada rígida em vez de regeneração da estampa.",
        size=11,
        bold=True,
        color=RED,
    )


def build() -> None:
    rows = read_rows()
    if len(rows) != 49:
        raise ValueError(f"Esperados 49 produtos; encontrados {len(rows)}")
    if any(not (CARDS_DIR / f"{row['product_id']}.jpg").exists() for row in rows):
        raise FileNotFoundError("Uma ou mais pranchas de produto estão ausentes")

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    page_break(doc)
    add_executive_summary(doc, rows)
    page_break(doc)
    add_methodology(doc)
    page_break(doc)
    add_priority_page(
        doc,
        rows,
        "REFAZER",
        "Prioridade 1 — refazer antes do lançamento",
        "Estes itens têm diferença visível de escala, inconsistência entre cores ou risco de a arte não ser fiel ao original.",
    )
    page_break(doc)
    add_priority_page(
        doc,
        rows,
        "REVISAR",
        "Prioridade 2 — revisar com atenção",
        "Estes itens estão próximos do aceitável, mas exigem conferência de escala, modelagem da peça ou microdetalhes antes de um aceite definitivo.",
    )
    for index, row in enumerate(rows, start=1):
        page_break(doc)
        add_card_page(doc, row, index, len(rows))
    page_break(doc)
    add_next_steps(doc)

    doc.core_properties.title = "Auditoria visual NIMBUS — dimensões de arte"
    doc.core_properties.subject = "Comparação YouDraw x fotos lifestyle da Nuvemshop"
    doc.core_properties.author = "NIMBUS"
    doc.core_properties.keywords = "NIMBUS, YouDraw, Nuvemshop, auditoria, escala, estampa"
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
